import os
from backend.app.services.parsers.base_parser import LocalParsedDocument
from backend.app.services.parsers.invoice_parser import InvoiceParser
from backend.app.services.parsers.packing_list_parser import PackingListParser
from backend.app.services.parsers.customs_docs_parser import CustomsDocsParser

class DefaultGenericParser:
    def parse(self, file_path: str, ext: str):
        documents = []
        try:
            # Processamento de arquivos PDF
            if ext == "pdf":
                import fitz  # PyMuPDF
                import pytesseract
                from PIL import Image
                import io

                doc = fitz.open(file_path)
                for i, page in enumerate(doc):
                    text = page.get_text("text")
                    if text.strip():
                        documents.append(LocalParsedDocument(text=text, page_number=i + 1))
                    else:
                        pix = page.get_pixmap()
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes))
                        ocr_text = pytesseract.image_to_string(img, lang="por+eng")
                        if ocr_text.strip():
                            documents.append(LocalParsedDocument(text=ocr_text, page_number=i + 1))

            # Processamento de imagens
            elif ext in ["png", "jpg", "jpeg"]:
                import pytesseract
                from PIL import Image

                img = Image.open(file_path)
                ocr_text = pytesseract.image_to_string(img, lang="por+eng")
                if ocr_text.strip():
                    documents.append(LocalParsedDocument(text=ocr_text, page_number=1))

            # Processamento de arquivos Word
            elif ext == "docx":
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                documents.append(LocalParsedDocument(text=text, page_number=1))

            else:
                documents.append(LocalParsedDocument(
                    text=f"[Conteúdo não extraído. Formato {ext} requer OCR ou biblioteca específica]", 
                    page_number=1
                ))

        except Exception as e:
            print(f"Erro ao processar {file_path}: {e}")
            
        if not documents:
             documents.append(LocalParsedDocument(text="[Documento vazio ou necessita de OCR]", page_number=1))
             
        return documents

class ParserServiceFactory:
    def __init__(self):
        # Nenhuma API Key é necessária para rodar localmente
        # Futuramente, podemos injetar os parsers especialistas dinamicamente
        self.specialized_parsers = [
            # Descomentar quando houver lógica de roteamento por conteúdo
            # InvoiceParser(),
            # PackingListParser(),
            # CustomsDocsParser()
        ]
        self.default_parser = DefaultGenericParser()

    async def parse_file(self, file_path: str):
        ext = file_path.lower().split('.')[-1]

        # Padrão Factory / Strategy:
        # Em uma implementação real avançada, leríamos as primeiras páginas para identificar se é Invoice ou Packing List.
        # Aqui deixamos a infraestrutura pronta.

        for parser in self.specialized_parsers:
            if parser.can_handle(ext):
                # O parser precisaria identificar se o conteúdo REALMENTE é dele.
                # Como não temos IA de classificação na factory ainda, cairemos para o default.
                pass

        return self.default_parser.parse(file_path, ext)

parser_service = ParserServiceFactory()
